import requests
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import os

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    if request.method == 'POST':
        erp = request.form['erp']
        processos = request.form.getlist('processos')
        informacoes_erp = request.form.getlist('informacoes_erp')

         # Formular a pergunta para o ChatGPT
        pergunta = (f"Qual seria a melhor análise funcional para realizar a integração do ERP {erp} com o nosso SaaS Paytrack? "
            f"Processos de integração desejados: {', '.join(processos)}. "
            f"Informações e campos necessários pelo ERP repassadas pelo cliente: {', '.join(informacoes_erp)}. "
            "Preciso que no retorno do documento contenha um mapeamento de campos olhando para o ERP selecionado, na linguagem do ERP "
            "Além de me retornar um JSON de exemplo formatado. "
            "É importante salientar, que o mapeamento de campos precisa vir no formato tabela e o JSON, "
            "E precisa ser formatado com as nomenclaturas do ERP, "
            "Como por exemplo bukrs que significa a empresa para o SAP e assim sucessivamente. "
            "Seguem algumas diretrizes que o nosso sistema Paytrack tem para integrar com os ERPs. "
            "1. - Utilizamos comunicação Sincrona com os Webservices do cliente. "
            "2. - A Paytrack é ativa nas integrações, ou seja, após este passo o cliente irá disponibilizar um Webservice para consumirmos. "
            "3. - A Análise funcional precisará ser separada no documento por cenário selecionado, ou seja, uma análise para adiantamento, uma para prestação de contas etc.")


        print("pergunta" + pergunta)
        headers = {'Authorization': f'Bearer {os.getenv("OPENAI_API_KEY")}'}
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [ {
                            "role": "assistant",
                            "content": "Você é um especialista em análise funcional de integrações, com conhecimento nos maiores ERPs de marcado com foco em documentações e mapeamentos"
                          },
                          {"role": "user",
                           "content": pergunta
                          }],
            "temperature": 0.7
        }

        response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=data)
        
        if response.status_code == 200:
            resposta_chatgpt = response.json()['choices'][0]['message']['content']

            # Início da criação do documento com os dados do formulário e a resposta do ChatGPT
            document = Document()

            # Adiciona o título com estilo personalizado diretamente
            title = document.add_heading(level=0)
            run = title.add_run('Integração ERP com Paytrack')
            run.font.size = Pt(14)
            run.font.bold = True
            
            # Adiciona as informações básicas
            paragraph = document.add_paragraph()
            paragraph.add_run(f"ERP: {erp}\n").bold = True
            paragraph.add_run(f"Processos desejados: {', '.join(processos)}\n").italic = True
            paragraph.add_run(f"Informações necessárias pelo ERP: {', '.join(informacoes_erp)}\n")

            # Adiciona a análise funcional recomendada
            document.add_heading('Análise Funcional Recomendada:', level=1)
            analysis_paragraph = document.add_paragraph(resposta_chatgpt)
            analysis_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Tabela de mapeamento de campos ERP
            document.add_page_break()
            document.add_heading('Mapeamento de Campos ERP:', level=1)
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Campo'
            hdr_cells[1].text = 'Descrição'
            for campo in informacoes_erp:
                row_cells = table.add_row().cells
                row_cells[0].text = campo
                row_cells[1].text = "Descrição para " + campo

            # Salvando e enviando o documento
            filename = "analise_integracao.docx"
            filepath = os.path.join(os.path.abspath(os.path.dirname(__file__)), filename)
            document.save(filepath)

            return send_file(filepath, as_attachment=True)
        else:
            print(response.json())  # Para ajudar na depuração
            return "Erro ao comunicar com a API do ChatGPT", 500

if __name__ == '__main__':
    app.run(debug=True)
